#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Apr 11 13:17:17 2018

@author: jayhan
"""

import docx 
doc = docx.Document() 
p1 = doc.add_heading('Hello world!')
p2 = doc.add_paragraph('paragraph 1')
doc.save('helloworld.docx')
