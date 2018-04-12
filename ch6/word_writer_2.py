#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Apr 11 13:17:17 2018

@author: jayhan
"""

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('dish.jpg', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

class record:
    def __init__(self,qty,_id,desc):
        self.qty = qty
        self.id = _id
        self.desc = desc
    
recordset = (record(1,100,'tomoto'),record(3,101,'lamb chop'),record(0.5,201,'sauce'))
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item.qty)
    row_cells[1].text = str(item.id)
    row_cells[2].text = item.desc

document.add_page_break()

document.save('demo.docx')